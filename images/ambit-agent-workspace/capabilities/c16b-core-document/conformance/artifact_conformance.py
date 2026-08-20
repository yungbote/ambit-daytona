from __future__ import annotations

import hashlib
import html
import json
import re
import stat
import sys
import warnings
import zipfile
from datetime import UTC, datetime
from html.parser import HTMLParser
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree

import mammoth
from docx import Document


FIXED_TIME = datetime(2024, 1, 1, 0, 0, 0, tzinfo=UTC)
FIXED_ZIP_TIME = (2024, 1, 1, 0, 0, 0)
OUTPUT = Path(sys.argv[1]).resolve()
ALLOWED_TAGS = frozenset(
    {
        "a",
        "blockquote",
        "br",
        "em",
        "h1",
        "h2",
        "h3",
        "h4",
        "li",
        "ol",
        "p",
        "strong",
        "table",
        "tbody",
        "td",
        "th",
        "thead",
        "tr",
        "ul",
    }
)
VOID_TAGS = frozenset({"br"})
ACTIVE_CONTAINERS = frozenset({"iframe", "math", "object", "script", "style", "svg"})
DROPPED_VOID_ELEMENTS = frozenset({"embed", "img", "input", "link", "meta", "source"})
MAMMOTH_STYLE_MAP = "p[style-name='Title'] => h1:fresh"


class SemanticHTMLSanitizer(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.output: list[str] = []
        self.open_tags: list[tuple[str, str]] = []
        self.rejected_elements = 0
        self.rejected_attributes = 0
        self.active_depth = 0
        self.table_row_counts: list[int] = []
        self.row_header_stack: list[bool] = []
        self.heading_count = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized = tag.lower()
        if self.active_depth:
            self.active_depth += 1
            return
        if normalized in ACTIVE_CONTAINERS:
            self.rejected_elements += 1
            self.active_depth = 1
            return
        if normalized in DROPPED_VOID_ELEMENTS:
            self.rejected_elements += 1
            self.rejected_attributes += len(attrs)
            return
        if normalized not in ALLOWED_TAGS:
            self.rejected_elements += 1
            self.open_tags.append((normalized, ""))
            return
        output_tag = normalized
        if normalized == "table":
            self.table_row_counts.append(0)
        elif normalized == "tr":
            is_header = bool(self.table_row_counts and self.table_row_counts[-1] == 0)
            self.row_header_stack.append(is_header)
            if self.table_row_counts:
                self.table_row_counts[-1] += 1
        elif normalized == "td" and self.row_header_stack and self.row_header_stack[-1]:
            output_tag = "th"
        admitted: list[tuple[str, str]] = []
        for name, value in attrs:
            attribute = name.lower()
            if normalized == "a" and attribute == "href" and value and value.startswith("#"):
                admitted.append(("href", value))
            elif output_tag == "th" and attribute == "scope" and value in {"col", "row"}:
                admitted.append(("scope", value))
            else:
                self.rejected_attributes += 1
        if output_tag == "th" and not any(name == "scope" for name, _ in admitted):
            admitted.append(("scope", "col"))
        if output_tag == "h1" and self.heading_count == 0:
            admitted.append(("id", "document-title"))
            self.heading_count += 1
        rendered_attributes = "".join(
            f' {name}="{html.escape(value, quote=True)}"' for name, value in sorted(admitted)
        )
        self.output.append(f"<{output_tag}{rendered_attributes}>")
        if normalized not in VOID_TAGS:
            self.open_tags.append((normalized, output_tag))

    def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized = tag.lower()
        if normalized in ACTIVE_CONTAINERS or normalized in DROPPED_VOID_ELEMENTS:
            self.rejected_elements += 1
            self.rejected_attributes += len(attrs)
            return
        self.handle_starttag(tag, attrs)
        if self.open_tags and self.open_tags[-1][0] == tag.lower():
            self.handle_endtag(tag)

    def handle_endtag(self, tag: str) -> None:
        if self.active_depth:
            self.active_depth -= 1
            return
        normalized = tag.lower()
        if normalized in DROPPED_VOID_ELEMENTS:
            return
        if normalized in VOID_TAGS:
            return
        if not self.open_tags or self.open_tags[-1][0] != normalized:
            raise ValueError(f"unbalanced Mammoth HTML tag: {normalized}")
        _, output_tag = self.open_tags.pop()
        if output_tag:
            self.output.append(f"</{output_tag}>")
        if normalized == "tr":
            self.row_header_stack.pop()
        elif normalized == "table":
            self.table_row_counts.pop()

    def handle_data(self, data: str) -> None:
        if not self.active_depth:
            self.output.append(html.escape(data, quote=False))

    def handle_entityref(self, name: str) -> None:
        if not self.active_depth:
            self.output.append(f"&{name};")

    def handle_charref(self, name: str) -> None:
        if not self.active_depth:
            self.output.append(f"&#{name};")

    def finish(self) -> str:
        self.close()
        if self.open_tags or self.active_depth or self.table_row_counts or self.row_header_stack:
            raise ValueError("unbalanced Mammoth HTML output")
        return "".join(self.output)


def sanitize_html(fragment: str) -> tuple[str, int, int]:
    sanitizer = SemanticHTMLSanitizer()
    sanitizer.feed(fragment)
    sanitized = sanitizer.finish()
    return sanitized, sanitizer.rejected_elements, sanitizer.rejected_attributes


def normalize_zip_archive(path: Path) -> None:
    normalized = path.with_name(f".{path.name}.normalized")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(normalized, "w") as target:
        for entry in sorted(source.infolist(), key=lambda item: item.filename):
            data = source.read(entry.filename)
            if entry.filename == "docProps/core.xml":
                for tag in (b"created", b"modified"):
                    pattern = rb"(<dcterms:" + tag + rb"[^>]*>)[^<]*(</dcterms:" + tag + rb">)"
                    data = re.sub(pattern, rb"\g<1>2024-01-01T00:00:00Z\g<2>", data)
            info = zipfile.ZipInfo(entry.filename, FIXED_ZIP_TIME)
            info.compress_type = entry.compress_type
            info.comment = entry.comment
            info.create_system = entry.create_system
            info.external_attr = entry.external_attr
            info.flag_bits = entry.flag_bits
            target.writestr(info, data, compress_type=entry.compress_type)
    normalized.replace(path)


def inspect_ooxml(path: Path, expected_last_paragraph: str) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        entries = archive.infolist()
        names = {entry.filename for entry in entries}
        assert len(names) == len(entries)
        assert "[Content_Types].xml" in names
        assert "word/document.xml" in names
        assert not any(name.lower().endswith("vbaproject.bin") for name in names)
        assert len(entries) < 512
        total_uncompressed = 0
        total_compressed = 0
        for entry in entries:
            pure = PurePosixPath(entry.filename)
            assert not pure.is_absolute()
            assert ".." not in pure.parts
            assert "\\" not in entry.filename
            assert not any(ord(character) <= 0x1F or ord(character) == 0x7F for character in entry.filename)
            assert entry.flag_bits & 0x1 == 0
            mode = (entry.external_attr >> 16) & 0xFFFF
            assert not stat.S_ISLNK(mode)
            assert entry.compress_type in {zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED}
            assert entry.file_size <= 16 * 1024 * 1024
            assert entry.compress_size <= 16 * 1024 * 1024
            if entry.file_size:
                assert entry.compress_size > 0
                assert entry.file_size / entry.compress_size <= 100
            total_uncompressed += entry.file_size
            total_compressed += entry.compress_size
        assert total_uncompressed <= 64 * 1024 * 1024
        assert total_compressed <= 32 * 1024 * 1024
        for xml_name in sorted(name for name in names if name.lower().endswith((".xml", ".rels"))):
            xml_payload = archive.read(xml_name)
            lowered = xml_payload.lower()
            assert b"<!doctype" not in lowered
            assert b"<!entity" not in lowered
        for relationship_name in sorted(name for name in names if name.lower().endswith(".rels")):
            root = ElementTree.fromstring(archive.read(relationship_name))
            for relationship in root:
                attributes = {name.lower(): value.lower() for name, value in relationship.attrib.items()}
                assert attributes.get("targetmode") != "external"
        content_types = archive.read("[Content_Types].xml").decode("utf-8")
        assert "macroenabled" not in content_types.lower()
    reopened = Document(path)
    assert reopened.paragraphs[0].text == "Ambit Project Brief"
    assert reopened.paragraphs[-1].text == expected_last_paragraph
    assert reopened.tables[0].cell(1, 1).text == "Passed"
    return {
        "entryCount": len(entries),
        "compressedBytes": total_compressed,
        "uncompressedBytes": total_uncompressed,
    }


def assert_unsafe_ooxml_rejected(
    source: Path,
    label: str,
    entries: list[zipfile.ZipInfo | tuple[str, bytes] | tuple[zipfile.ZipInfo, bytes]],
) -> None:
    fixture = OUTPUT / "document" / f"unsafe-{label}.docx"
    with zipfile.ZipFile(source) as original, zipfile.ZipFile(fixture, "w") as target:
        for item in original.infolist():
            target.writestr(item, original.read(item.filename), compress_type=item.compress_type)
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            for entry in entries:
                if isinstance(entry, tuple):
                    name_or_info, payload = entry
                    target.writestr(name_or_info, payload)
                else:
                    target.writestr(entry, b"unsafe")
    try:
        inspect_ooxml(fixture, "never reached")
    except (AssertionError, ElementTree.ParseError, ValueError, zipfile.BadZipFile):
        pass
    else:
        raise AssertionError(f"unsafe OOXML fixture was accepted: {label}")
    finally:
        fixture.unlink()


def exercise_unsafe_ooxml_fixtures(source: Path) -> list[str]:
    duplicate = ("word/document.xml", b"duplicate")
    symlink = zipfile.ZipInfo("word/unsafe-link", FIXED_ZIP_TIME)
    symlink.create_system = 3
    symlink.external_attr = (stat.S_IFLNK | 0o777) << 16
    unsupported = zipfile.ZipInfo("word/unsupported.bin", FIXED_ZIP_TIME)
    unsupported.compress_type = zipfile.ZIP_BZIP2
    high_ratio = zipfile.ZipInfo("word/high-ratio.bin", FIXED_ZIP_TIME)
    high_ratio.compress_type = zipfile.ZIP_DEFLATED
    cases: list[
        tuple[str, list[zipfile.ZipInfo | tuple[str, bytes] | tuple[zipfile.ZipInfo, bytes]]]
    ] = [
        ("duplicate-name", [duplicate]),
        ("symlink-entry", [symlink]),
        ("backslash-path", [("word\\escape.xml", b"<safe/>")]),
        ("control-path", [("word/control\x01.xml", b"<safe/>")]),
        ("unsupported-compression", [unsupported]),
        ("expansion-ratio", [(high_ratio, b"0" * (1024 * 1024))]),
        ("doctype-entity", [("word/unsafe.xml", b'<!DOCTYPE x [<!ENTITY e "unsafe">]><x>&e;</x>')]),
        ("case-insensitive-macro", [("word/VBAPROJECT.BIN", b"macro")]),
        (
            "external-relationship",
            [
                (
                    "custom/_rels/unsafe.RELS",
                    b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                    b'<Relationship Id="rId1" Type="unsafe" Target="https://example.com" TargetMode="EXTERNAL"/>'
                    b"</Relationships>",
                )
            ],
        ),
    ]
    for label, entries in cases:
        assert_unsafe_ooxml_rejected(source, label, entries)
    encrypted_fixture = OUTPUT / "document" / "unsafe-encrypted-entry.docx"
    encrypted_payload = bytearray(source.read_bytes())
    local_offset = encrypted_payload.find(b"PK\x03\x04")
    central_offset = encrypted_payload.find(b"PK\x01\x02")
    assert local_offset >= 0 and central_offset >= 0
    local_flags = int.from_bytes(encrypted_payload[local_offset + 6 : local_offset + 8], "little") | 0x1
    central_flags = int.from_bytes(encrypted_payload[central_offset + 8 : central_offset + 10], "little") | 0x1
    encrypted_payload[local_offset + 6 : local_offset + 8] = local_flags.to_bytes(2, "little")
    encrypted_payload[central_offset + 8 : central_offset + 10] = central_flags.to_bytes(2, "little")
    encrypted_fixture.write_bytes(encrypted_payload)
    try:
        inspect_ooxml(encrypted_fixture, "never reached")
    except (AssertionError, RuntimeError, zipfile.BadZipFile):
        pass
    else:
        raise AssertionError("unsafe OOXML fixture was accepted: encrypted-entry")
    finally:
        encrypted_fixture.unlink()
    return [label for label, _ in cases] + ["encrypted-entry"]


def create_document() -> tuple[Path, dict[str, int]]:
    path = OUTPUT / "document" / "project-brief-v1.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Ambit Project Brief", level=0)
    document.add_paragraph("A minimal, structurally validated core document workflow.")
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Control"
    table.rows[0].cells[1].text = "Result"
    for control, result in (
        ("Non-root runtime", "Passed"),
        ("Offline execution", "Passed"),
        ("Exact pack digest", "Bound"),
    ):
        cells = table.add_row().cells
        cells[0].text = control
        cells[1].text = result
    conclusion = "Conclusion: the document remains editable and structurally inspectable."
    document.add_paragraph(conclusion)
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME
    document.save(path)
    normalize_zip_archive(path)
    return path, inspect_ooxml(path, conclusion)


def revise_document(source: Path) -> tuple[Path, dict[str, int]]:
    original_digest = hashlib.sha256(source.read_bytes()).hexdigest()
    revised = OUTPUT / "document" / "project-brief-v2.docx"
    document = Document(source)
    document.add_heading("Verified repair", level=1)
    conclusion = "The second revision preserves the original structure and adds a reviewed conclusion."
    document.add_paragraph(conclusion)
    document.core_properties.modified = FIXED_TIME
    document.save(revised)
    normalize_zip_archive(revised)
    assert hashlib.sha256(source.read_bytes()).hexdigest() == original_digest
    return revised, inspect_ooxml(revised, conclusion)


def derive_html_preview(source: Path, revision: str, expected_text: tuple[str, ...]) -> tuple[Path, dict[str, object]]:
    with source.open("rb") as stream:
        converted = mammoth.convert_to_html(stream, style_map=MAMMOTH_STYLE_MAP)
    sanitized, rejected_elements, rejected_attributes = sanitize_html(converted.value)
    with source.open("rb") as stream:
        repeated = mammoth.convert_to_html(stream, style_map=MAMMOTH_STYLE_MAP)
    repeated_sanitized, repeated_elements, repeated_attributes = sanitize_html(repeated.value)
    assert repeated_sanitized == sanitized
    assert repeated_elements == rejected_elements
    assert repeated_attributes == rejected_attributes
    assert [str(message) for message in repeated.messages] == [str(message) for message in converted.messages]
    assert rejected_elements == 0
    assert rejected_attributes == 0
    assert sanitized.startswith('<h1 id="document-title">Ambit Project Brief</h1>')
    normalized_text = " ".join(re.sub(r"<[^>]+>", " ", sanitized).split())
    for value in expected_text:
        assert " ".join(value.split()) in normalized_text
    assert not re.search(r"\son[a-z]+=", sanitized, flags=re.IGNORECASE)
    preview = OUTPUT / "document" / f"project-brief-{revision}.derived.html"
    preview.write_text(
        "<!doctype html>\n"
        '<html lang="en"><head><meta charset="utf-8">'
        '<meta name="ambit-preview" content="derived-non-layout-authoritative">'
        f'<title>Ambit Project Brief {revision}</title></head><body>'
        f'<main aria-labelledby="document-title">{sanitized}</main>'
        "</body></html>\n"
    )
    return preview, {
        "conversionMessages": [str(message) for message in converted.messages],
        "rejectedElements": rejected_elements,
        "rejectedAttributes": rejected_attributes,
    }


def digest(path: Path) -> dict[str, object]:
    payload = path.read_bytes()
    return {
        "path": str(path.relative_to(OUTPUT)),
        "bytes": len(payload),
        "sha256": hashlib.sha256(payload).hexdigest(),
    }


OUTPUT.mkdir(parents=True, exist_ok=True)
document_v1, inspection_v1 = create_document()
unsafe_ooxml_cases = exercise_unsafe_ooxml_fixtures(document_v1)
preview_v1, preview_receipt_v1 = derive_html_preview(
    document_v1,
    "v1",
    ("Ambit Project Brief", "Non-root runtime", "Passed"),
)
document_v2, inspection_v2 = revise_document(document_v1)
preview_v2, preview_receipt_v2 = derive_html_preview(
    document_v2,
    "v2",
    ("Ambit Project Brief", "Verified repair", "reviewed conclusion"),
)

malicious, rejected_elements, rejected_attributes = sanitize_html(
    '<p onclick="steal()">safe<script>alert(1)</script>'
    '<a href="https://example.com" onmouseover="steal()">link</a></p>'
    '<section>benign https:// javascript: &lt;script&gt;</section>'
    '<img src="https://example.com/tracker.png" onerror="steal()">'
    '<a href="#local">internal</a>'
)
assert malicious == '<p>safe<a>link</a></p>benign https:// javascript: &lt;script&gt;<a href="#local">internal</a>'
assert rejected_elements == 3
assert rejected_attributes == 5

artifacts = [document_v1, preview_v1, document_v2, preview_v2]
(OUTPUT / "artifact-receipt.json").write_text(
    json.dumps(
        {
            "schema": "ambit.runtime-pack-document-conformance/v3",
            "artifacts": sorted((digest(path) for path in artifacts), key=lambda item: str(item["path"])),
            "document": {
                "workflow": ["create", "inspect", "derive-html-preview", "edit", "inspect", "derive-html-preview", "commit-candidate"],
                "originalPreserved": True,
                "revisionCount": 2,
                "macroExecution": "disabled",
                "macroPayload": "absent",
                "externalRelationships": "denied",
                "ooxmlInspection": {"v1": inspection_v1, "v2": inspection_v2},
                "rejectedUnsafeOOXMLCases": unsafe_ooxml_cases,
            },
            "preview": {
                "format": "sanitized-semantic-html",
                "derived": True,
                "layoutAuthoritative": False,
                "nativeRenderFidelityOwner": "C19.ArtifactRenderValidatePort",
                "externalFetches": "denied",
                "activeContent": "denied",
                "v1": preview_receipt_v1,
                "v2": preview_receipt_v2,
            },
            "movedToSpecialistOwners": {
                "C18": [
                    "spreadsheet",
                    "presentation",
                    "pdf-specialist-and-ocr",
                    "data-analysis",
                    "research-and-publishing",
                    "web-application-browser",
                    "media-and-diagrams",
                    "node-typescript-runtime",
                    "language-intelligence-and-lsp",
                ],
                "C19": [
                    "native-document-render-selection",
                    "layout-and-font-fidelity",
                    "page-rasterization-and-visual-validation",
                ],
            },
        },
        indent=2,
        sort_keys=True,
    )
    + "\n"
)
